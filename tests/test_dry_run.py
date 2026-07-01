from __future__ import annotations

import json
import io
import sys
import tempfile
import unittest
from datetime import date
from pathlib import Path
from unittest.mock import patch
from contextlib import redirect_stdout

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from recruit_crawler.config import ConfigError, load_config
from recruit_crawler.cli import main as cli_main
from recruit_crawler.capture_import import build_capture_quality_gate, import_capture_files, select_capture_files
from recruit_crawler.jd_parser import parse_deadline
from recruit_crawler.pipeline import run_capture_import, run_dry_run, run_live_run
from recruit_crawler.browser_evidence import build_browser_evidence, _redact
from recruit_crawler.relevance import evaluate_relevance_cases
from recruit_crawler.scorer import score_snapshot
from recruit_crawler.schemas import JDSnapshot, RelevanceCase, UserContext
from recruit_crawler.user_context import (
    UserContextImportError,
    merge_supplemental_answers,
    parse_context_document,
    supplemental_questions,
)
from recruit_crawler.schemas import SourceManifest
from recruit_crawler.sources.base import build_source_adapter
from recruit_crawler.sources.http import (
    HttpResponse,
    PublicJobsHttpAdapter,
    SourceAccessError,
    _SafeRedirectHandler,
)
from recruit_crawler.sources.platforms import (
    JobKoreaAdapter,
    JumpitAdapter,
    RallitAdapter,
    RocketPunchBrowserAutomationAdapter,
    LinkedInAdapter,
    SaraminAdapter,
    WantedAdapter,
)


CONFIG = ROOT / "config" / "sample_config.json"


class DryRunTests(unittest.TestCase):
    def test_fixture_e2e_generates_report_without_expired_postings(self) -> None:
        config = load_config(CONFIG)
        summary, report, ranked = run_dry_run(config, date(2026, 6, 30))

        self.assertTrue(summary.report_path.exists())
        self.assertEqual(summary.candidates_collected, 6)
        self.assertEqual(summary.duplicates_removed, 1)
        self.assertEqual(summary.expired_excluded, 1)
        self.assertEqual(summary.ranked_count, 4)
        self.assertIn("# 오늘의 채용 후보", report)
        self.assertIn("## 우선순위 표", report)
        self.assertIn("## 상세 메모", report)
        self.assertIn("https://jobs.example.test/apply-ml-engineer", report)
        self.assertNotIn("Expired ML Intern", report)
        self.assertNotIn("RAW_JD_CANARY_EXPIRED", report)
        self.assertEqual(len(ranked), 4)

    def test_recommendation_buckets_include_apply_hold_and_low_priority(self) -> None:
        config = load_config(CONFIG)
        _summary, _report, ranked = run_dry_run(config, date(2026, 6, 30))
        recommendations = {item.recommendation for item in ranked}

        self.assertIn("apply", recommendations)
        self.assertIn("hold", recommendations)
        self.assertIn("low_priority", recommendations)

    def test_report_excludes_raw_jd_and_private_profile_canaries(self) -> None:
        config = load_config(CONFIG)
        _summary, report, _ranked = run_dry_run(config, date(2026, 6, 30))

        forbidden = [
            "RAW_JD_CANARY_APPLY",
            "RAW_JD_CANARY_HOLD",
            "RAW_JD_CANARY_LOW",
            "RAW_JD_CANARY_AMBIGUOUS",
            "RAW_JD_CANARY_DUPLICATE",
            "PRIVATE_PROFILE_CANARY",
            "Ignore previous instructions",
        ]
        for value in forbidden:
            self.assertNotIn(value, report)

    def test_each_selected_posting_has_actionable_report_fields(self) -> None:
        config = load_config(CONFIG)
        _summary, report, ranked = run_dry_run(config, date(2026, 6, 30))

        self.assertEqual(len(ranked), config.top_n)
        for assessment in ranked:
            snapshot = assessment.snapshot
            self.assertIn(snapshot.source_url, report)
            self.assertIn(f"(`{assessment.recommendation}`)", report)
            self.assertIn(f"점수 **{assessment.score}**", report)
            self.assertIn("| 항목 | 내용 |", report)
            self.assertIn("- **맞는 부분**:", report)
            self.assertIn("- **리스크**:", report)
            self.assertIn("- **확인할 것**:", report)
            self.assertIn("- **지원 각도**:", report)

    def test_report_surface_text_is_korean(self) -> None:
        config = load_config(CONFIG)
        _summary, report, _ranked = run_dry_run(config, date(2026, 6, 30))

        english_labels = [
            "Recruiting Dry-Run Report",
            "Run date",
            "Top Candidates",
            "Recommendation:",
            "Estimated fit score",
            "Structured snapshot",
            "Matched evidence",
            "Verification questions",
            "Positioning seed",
            "No major structured risk detected",
        ]
        for label in english_labels:
            self.assertNotIn(label, report)

    def test_dry_run_does_not_reference_personal_info_inputs(self) -> None:
        code_files = [
            ROOT / "src" / "recruit_crawler" / "pipeline.py",
            ROOT / "src" / "recruit_crawler" / "sources" / "fixture.py",
            ROOT / "src" / "recruit_crawler" / "sources" / "base.py",
            ROOT / "src" / "recruit_crawler" / "summarizer.py",
            ROOT / "src" / "recruit_crawler" / "report_writer.py",
        ]
        for path in code_files:
            self.assertNotIn("personal_info", path.read_text(encoding="utf-8"))

    def test_unknown_deadline_is_uncertain_not_expired(self) -> None:
        parsed, uncertain = parse_deadline("not listed")

        self.assertIsNone(parsed)
        self.assertTrue(uncertain)

    def test_top_n_is_configurable(self) -> None:
        raw = json.loads(CONFIG.read_text(encoding="utf-8"))
        raw["top_n"] = 2
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            config_dir = tmp_path / "config"
            config_dir.mkdir()
            config_path = config_dir / "sample_config.json"
            raw["fixture_path"] = str(ROOT / "fixtures" / "postings.json")
            raw["output_dir"] = str(tmp_path / "reports")
            config_path.write_text(json.dumps(raw), encoding="utf-8")

            config = load_config(config_path)
            summary, _report, ranked = run_dry_run(config, date(2026, 6, 30))

        self.assertEqual(summary.ranked_count, 2)
        self.assertEqual(len(ranked), 2)

    def test_real_source_is_blocked_for_dry_run(self) -> None:
        raw = json.loads(CONFIG.read_text(encoding="utf-8"))
        raw["sources"].append(
            {
                "source_id": "real_example",
                "enabled": True,
                "access_mode": "public_page",
                "auth_required": False,
                "tos_review_status": "unknown",
                "domains": ["example.com"],
                "rate_limit": "unknown",
                "failure_mode": "skip_source",
                "allowed_persisted_fields": [],
            }
        )
        with tempfile.TemporaryDirectory() as tmp:
            config_path = Path(tmp) / "config.json"
            config_path.write_text(json.dumps(raw), encoding="utf-8")
            with self.assertRaises(ConfigError):
                load_config(config_path)

    def test_dry_run_rejects_preloaded_real_source_config(self) -> None:
        config = load_config(ROOT / "config" / "live_sources.sample.json", allow_real_sources=True)

        with self.assertRaises(ConfigError):
            run_dry_run(config, date(2026, 6, 30))

    def test_manual_local_source_mode_is_allowed_without_real_adapter_enablement(self) -> None:
        raw = json.loads(CONFIG.read_text(encoding="utf-8"))
        raw["sources"][0]["access_mode"] = "manual"
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            config_path = tmp_path / "config.json"
            raw["fixture_path"] = str(ROOT / "fixtures" / "postings.json")
            raw["output_dir"] = str(tmp_path / "reports")
            config_path.write_text(json.dumps(raw), encoding="utf-8")
            config = load_config(config_path)
            summary, _report, ranked = run_dry_run(config, date(2026, 6, 30))

        self.assertEqual(summary.sources_attempted, ["fixture"])
        self.assertEqual(len(ranked), raw["top_n"])

    def test_live_config_can_load_reviewed_real_sources(self) -> None:
        raw = json.loads(CONFIG.read_text(encoding="utf-8"))
        raw["sources"] = [
            {
                "source_id": "company_careers",
                "enabled": True,
                "access_mode": "public_page",
                "auth_required": False,
                "tos_review_status": "pass",
                "domains": ["example.com"],
                "rate_limit": "1 request / second",
                "failure_mode": "skip_source",
                "allowed_persisted_fields": [],
                "options": {"start_urls": ["https://example.com/careers/ml"], "require_robots": False},
            }
        ]
        with tempfile.TemporaryDirectory() as tmp:
            config_path = Path(tmp) / "config.json"
            config_path.write_text(json.dumps(raw), encoding="utf-8")
            config = load_config(config_path, allow_real_sources=True)

        self.assertEqual(config.sources[0].source_id, "company_careers")

    def test_public_jobs_adapter_extracts_json_ld_job_posting(self) -> None:
        manifest = SourceManifest(
            source_id="company_careers",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["example.com"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "start_urls": ["https://example.com/jobs/ml"],
                "require_robots": False,
                "explicit_automated_permission": True,
                "delay_seconds": 0,
            },
        )
        html = """
        <html><head>
          <script type="application/ld+json">
          {
            "@type": "JobPosting",
            "title": "Machine Learning Engineer",
            "url": "https://example.com/jobs/ml",
            "validThrough": "2026-08-01T00:00:00+09:00",
            "hiringOrganization": {"name": "Example AI"},
            "jobLocation": {"address": {"addressLocality": "Seoul"}},
            "qualifications": "Python and machine learning",
            "responsibilities": "Build ranking models"
          }
          </script>
        </head></html>
        """
        adapter = PublicJobsHttpAdapter(manifest)
        with patch.object(adapter, "_fetch", return_value=HttpResponse("https://example.com/jobs/ml", html)):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertEqual(candidates[0].title, "Machine Learning Engineer")
        self.assertEqual(candidates[0].company, "Example AI")
        self.assertEqual(candidates[0].deadline_raw, "2026-08-01")

    def test_public_jobs_adapter_filters_links_and_candidates_by_keywords(self) -> None:
        manifest = SourceManifest(
            source_id="jobkorea",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.jobkorea.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "search_urls": ["https://www.jobkorea.co.kr/recruit/ai-jobs"],
                "include_url_patterns": ["/Recruit/GI_Read"],
                "link_include_keywords": ["python", "데이터"],
                "candidate_include_keywords": ["python", "데이터"],
                "candidate_exclude_keywords": ["주방"],
                "require_robots": False,
                "explicit_automated_permission": True,
                "delay_seconds": 0,
            },
        )
        listing_html = """
        <a href="/Recruit/GI_Read/1">Python 데이터 엔지니어</a>
        <a href="/Recruit/GI_Read/2">주방 보조 모집</a>
        """
        job_html = """
        <html><head><title>Python 데이터 엔지니어 - Example</title>
        <meta name="description" content="Python과 데이터 파이프라인 업무"></head></html>
        """
        adapter = PublicJobsHttpAdapter(manifest)

        def fake_fetch(url: str) -> HttpResponse:
            if url.endswith("/recruit/ai-jobs"):
                return HttpResponse(url, listing_html)
            return HttpResponse(url, job_html)

        with patch.object(adapter, "_fetch", side_effect=fake_fetch):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertIn("Python", candidates[0].title)

    def test_http_adapter_blocks_off_domain_redirect_before_following(self) -> None:
        manifest = SourceManifest(
            source_id="company_careers",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["example.com"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={},
        )
        adapter = PublicJobsHttpAdapter(manifest)
        handler = _SafeRedirectHandler(adapter)

        with self.assertRaises(SourceAccessError):
            handler.redirect_request(
                type("Req", (), {"full_url": "https://example.com/jobs/ml"})(),
                None,
                302,
                "Found",
                {},
                "https://evil.example/jobs/ml",
            )

    def test_linkedin_adapter_requires_explicit_approved_access(self) -> None:
        manifest = SourceManifest(
            source_id="linkedin",
            enabled=True,
            access_mode="api",
            auth_required=True,
            tos_review_status="pass",
            domains=["www.linkedin.com"],
            rate_limit="approved partner/API access only",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={"approved_partner_access": False},
        )
        adapter = PublicJobsHttpAdapter(manifest)

        with self.assertRaises(SourceAccessError):
            adapter.collect()

    def test_known_platforms_use_platform_specific_adapters(self) -> None:
        expected = {
            "jumpit": JumpitAdapter,
            "saramin": SaraminAdapter,
            "jobkorea": JobKoreaAdapter,
            "wanted": WantedAdapter,
            "linkedin": LinkedInAdapter,
            "rallit": RallitAdapter,
            "rocketpunch": RocketPunchBrowserAutomationAdapter,
        }
        for source_id, adapter_class in expected.items():
            manifest = SourceManifest(
                source_id=source_id,
                enabled=True,
                access_mode="public_page" if source_id != "linkedin" else "api",
                auth_required=source_id == "linkedin",
                tos_review_status="pass",
                domains=["example.com"],
                rate_limit="1 request / second",
                failure_mode="skip_source",
                allowed_persisted_fields=[],
                options={},
            )

            adapter = build_source_adapter(manifest, ROOT / "fixtures" / "postings.json")

            self.assertIsInstance(adapter, adapter_class)

    def test_jumpit_adapter_collects_jobs_from_sitemap_and_react_data(self) -> None:
        manifest = SourceManifest(
            source_id="jumpit",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["jumpit.saramin.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "sitemap_urls": ["https://jumpit.saramin.co.kr/sitemap/sitemap_position_view_1.xml"],
                "candidate_include_keywords": ["python"],
                "max_pages": 5,
                "require_robots": False,
                "explicit_automated_permission": True,
                "delay_seconds": 0,
            },
        )
        sitemap_xml = """
        <urlset><url><loc>https://jumpit.saramin.co.kr/position/54308479</loc></url></urlset>
        """
        page_html = r'''
        <script>
        self.__next_f.push([1,"{\"data\":{\"NOTICE\":[{\"title\":\"[당첨자 발표] 훈련과정 인증 이벤트\"}]}}"])
        self.__next_f.push([1,"{\"data\":{\"id\":54308479,\"title\":\"AI Platform Engineer\",\"companyName\":\"Example AI\",\"techStacks\":[{\"stack\":\"Python\"},{\"stack\":\"PyTorch\"}],\"responsibility\":\"Build ML serving systems\",\"qualifications\":\"Python and machine learning\",\"preferredRequirements\":\"LLM experience\",\"serviceInfo\":\"AI products\",\"newcomer\":false,\"minCareer\":2,\"closedAt\":\"2026-07-29 23:59:59\",\"location\":\"서울 강남구\"},\"dataUpdateCount\":1}"])
        </script>
        '''
        adapter = JumpitAdapter(manifest)

        def fake_fetch(url: str) -> HttpResponse:
            if url.endswith("sitemap_position_view_1.xml"):
                return HttpResponse(url, sitemap_xml)
            return HttpResponse(url, page_html)

        with patch.object(adapter, "_fetch", side_effect=fake_fetch):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertEqual(candidates[0].title, "AI Platform Engineer")
        self.assertEqual(candidates[0].company, "Example AI")
        self.assertEqual(candidates[0].deadline_raw, "2026-07-29")
        self.assertEqual(candidates[0].location, "서울 강남구")
        self.assertIn("Python", candidates[0].raw_jd["required_qualifications"])
        self.assertEqual(candidates[0].raw_jd["experience_tags"], ["경력2년↑"])

    def test_jumpit_adapter_tolerates_spaced_next_payload_shape(self) -> None:
        manifest = SourceManifest(
            source_id="jumpit",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["jumpit.saramin.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "sitemap_urls": ["https://jumpit.saramin.co.kr/sitemap/sitemap_position_view_1.xml"],
                "candidate_include_keywords": ["python"],
                "require_robots": False,
                "explicit_automated_permission": True,
                "delay_seconds": 0,
            },
        )
        sitemap_xml = "<urlset><url><loc>https://jumpit.saramin.co.kr/position/54308479</loc></url></urlset>"
        page_html = '''
        <script>
        self.__next_f.push([1, '{"data": {"title": "AI Platform Engineer", "companyName": "Example AI",
          "techStacks": [{"stack": "Python"}, {"stack": "PyTorch"}],
          "responsibility": "Build Python services", "qualifications": "Python ML",
          "preferredRequirements": "LLM", "serviceInfo": "AI products",
          "newcomer": true, "minCareer": 1, "closedAt": "2026-07-29 23:59:59",
          "location": "서울 강남구"}}'])
        </script>
        '''
        adapter = JumpitAdapter(manifest)

        with patch.object(
            adapter,
            "_fetch",
            side_effect=[
                HttpResponse("https://jumpit.saramin.co.kr/sitemap/sitemap_position_view_1.xml", sitemap_xml),
                HttpResponse("https://jumpit.saramin.co.kr/position/54308479", page_html),
            ],
        ):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertEqual(candidates[0].title, "AI Platform Engineer")
        self.assertIn("PyTorch", candidates[0].raw_jd["required_qualifications"])
        self.assertEqual(candidates[0].raw_jd["experience_tags"], ["신입", "경력1년↑"])

    def test_rallit_adapter_collects_jobs_from_public_positions(self) -> None:
        manifest = SourceManifest(
            source_id="rallit",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.rallit.com"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "search_urls": ["https://www.rallit.com"],
                "candidate_include_keywords": ["python"],
                "max_pages": 5,
                "require_robots": False,
                "explicit_automated_permission": True,
                "delay_seconds": 0,
            },
        )
        list_html = """
        <a href="/positions/1797/%EC%A3%BC%EC%8B%9D%ED%9A%8C%EC%82%AC-%EC%9C%A0%EB%8B%88%EC%9C%A0%EB%8B%88-ai-%EA%B0%9C%EB%B0%9C%EC%9E%90-%EC%B1%84%EC%9A%A9">
          주식회사 유니유니 AI 개발자 Python
        </a>
        """
        detail_html = """
        <html><head>
          <meta property="og:title" content="주식회사 유니유니 [주식회사 유니유니] ‘AI 개발자’ 채용 채용 - 랠릿" />
        </head><body>
          <h1>[주식회사 유니유니] ‘AI 개발자’ 채용</h1>
          <section>주식회사 유니유니, 어떤 곳인가요? 프라이버시 테크 AI 스타트업입니다.</section>
          <section>[주식회사 유니유니] ‘AI 개발자’ 채용, 어떤 일을 하나요?</section>
          #인공지능(AI) #컴퓨터 비전 #Python
          <h3>주요업무</h3>
          <p>SAVVY 서비스에 사용중인 딥러닝 모델의 개선 방안 연구 및 테스트 진행</p>
          <h3>자격요건</h3>
          <p>Python을 이용한 개발 경험 및 PyTorch 숙련도를 보유하신 분</p>
          <h3>우대사항</h3>
          <p>컴퓨터 비전 분야 실무 경험</p>
          <h3>혜택 및 복지</h3>
          <p>교육비 지원</p>
          <p>근무 지역 경기 성남시 분당구 판교로289번길 20</p>
          <p>경력 미들 (4~8년)</p>
          <p>최소 연봉 회사 내규에 따름</p>
          <p>마감일 채용 시 마감</p>
          <p>회사명 주식회사 유니유니 7 지원하기</p>
        </body></html>
        """
        adapter = RallitAdapter(manifest)

        def fake_fetch(url: str) -> HttpResponse:
            if url == "https://www.rallit.com":
                return HttpResponse(url, list_html)
            return HttpResponse(url, detail_html)

        with patch.object(adapter, "_fetch", side_effect=fake_fetch):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        candidate = candidates[0]
        self.assertEqual(candidate.source_id, "rallit")
        self.assertEqual(candidate.source_posting_id, "1797")
        self.assertEqual(candidate.title, "[주식회사 유니유니] ‘AI 개발자’ 채용")
        self.assertEqual(candidate.company, "주식회사 유니유니")
        self.assertEqual(candidate.deadline_raw, "채용시")
        self.assertIn("경기 성남시", candidate.location)
        self.assertIn("Python", candidate.raw_jd["required_qualifications"])
        self.assertIn("딥러닝 모델", candidate.raw_jd["responsibilities"][0])
        self.assertEqual(candidate.raw_jd["experience_tags"], ["미들 (4~8년)"])

    def test_rallit_adapter_strips_embedded_css_from_public_detail(self) -> None:
        manifest = SourceManifest(
            source_id="rallit",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.rallit.com"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "search_urls": ["https://www.rallit.com"],
                "candidate_include_keywords": ["python"],
                "require_robots": False,
                "explicit_automated_permission": True,
                "delay_seconds": 0,
            },
        )
        list_html = '<a href="/positions/3333/python-backend">Python 백엔드</a>'
        detail_html = """
        <html><body>
          <style>.mantine-9qoqdi{display:flex;color:#00CCAA;}@keyframes animation-1a410py{from{opacity:0.6;}to{opacity:0;}}</style>
          <h1>Python 백엔드 개발자</h1>
          <section>누아, 어떤 곳인가요? 인공지능 여행 기술 기업입니다.</section>
          <h3>주요업무</h3>
          <p>예약 관리 시스템 개발</p>
          <h3>자격요건</h3>
          <p>Python API 개발 경험</p>
          <h3>우대사항</h3>
          <p>여행 플랫폼 연동 경험</p>
          <p>회사명 누아 1 지원하기</p>
        </body></html>
        """
        adapter = RallitAdapter(manifest)

        with patch.object(
            adapter,
            "_fetch",
            side_effect=[
                HttpResponse("https://www.rallit.com", list_html),
                HttpResponse("https://www.rallit.com/positions/3333/python-backend", detail_html),
            ],
        ):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        serialized = json.dumps(candidates[0].raw_jd, ensure_ascii=False)
        self.assertIn("Python API 개발", serialized)
        self.assertNotIn("mantine", serialized)
        self.assertNotIn("keyframes", serialized)
        self.assertNotIn("display:flex", serialized)

    def test_rallit_adapter_tolerates_alternate_section_markers(self) -> None:
        manifest = SourceManifest(
            source_id="rallit",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.rallit.com"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "search_urls": ["https://www.rallit.com"],
                "candidate_include_keywords": ["python"],
                "require_robots": False,
                "explicit_automated_permission": True,
                "delay_seconds": 0,
            },
        )
        list_html = '<a href="/positions/4444/python-platform">Python 플랫폼</a>'
        detail_html = """
        <html><head><meta property="og:title" content="테크랩 [테크랩] Python 플랫폼 개발자 채용 - 랠릿" /></head><body>
          <h1>[테크랩] Python 플랫폼 개발자</h1>
          <section>테크랩, 어떤 곳인가요? AI 플랫폼 기업입니다.</section>
          <h3>합류하면 하게 될 업무</h3>
          <p>Python 기반 데이터 플랫폼을 개발합니다.</p>
          <h3>지원자격</h3>
          <p>Python API와 SQL 경험</p>
          <h3>우대사항</h3>
          <p>LLM 서비스 경험</p>
          <p>근무 지역 서울 강남구</p>
          <p>경력 주니어 (1~3년)</p>
          <p>마감일 2026.07.31</p>
          <p>회사명 테크랩 3 지원하기</p>
        </body></html>
        """
        adapter = RallitAdapter(manifest)

        with patch.object(
            adapter,
            "_fetch",
            side_effect=[
                HttpResponse("https://www.rallit.com", list_html),
                HttpResponse("https://www.rallit.com/positions/4444/python-platform", detail_html),
            ],
        ):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertIn("데이터 플랫폼", candidates[0].raw_jd["responsibilities"][0])
        self.assertIn("Python API", candidates[0].raw_jd["required_qualifications"][0])
        self.assertEqual(candidates[0].deadline_raw, "2026-07-31")

    def test_rocketpunch_browser_automation_parses_listing_cards_without_detail_links(self) -> None:
        listing_url = "https://www.rocketpunch.com/en/jobs"
        manifest = SourceManifest(
            source_id="rocketpunch",
            enabled=True,
            access_mode="browser_automation",
            auth_required=False,
            tos_review_status="unknown",
            domains=["www.rocketpunch.com"],
            rate_limit="browser automation with user-directed notice override",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "search_urls": [listing_url],
                "candidate_include_keywords": ["python", "data", "backend", "fintech"],
                "policy_override_mode": "user_directed_ignore",
                "policy_override_reason": "User directed RocketPunch browser automation despite source notice.",
                "policy_override_acknowledges_source_notice": True,
                "max_pages": 10,
                "delay_seconds": 0,
                "fetch_detail_pages": False,
            },
        )
        dom_html = """
        <html><body>
          <div class="listing-card">
            <a href="/en/jobs/158927?list=true">detail</a>
            <span class="company-name">페이데이터</span>
            <h2 class="job-title">DataOps Engineer</h2>
            <p>Python SQL Airflow data pipeline operations</p>
            <span>Seoul</span>
          </div>
          <div class="listing-card">
            <span class="company-name">핀테크랩</span>
            <h2 class="job-title">Fintech Project Manager</h2>
            <p>Fintech AI backend platform coordination</p>
          </div>
          <div class="listing-card">
            <span class="company-name">디자인랩</span>
            <h2 class="job-title">Product Designer</h2>
            <p>Design system and marketing assets</p>
          </div>
        </body></html>
        """
        adapter = RocketPunchBrowserAutomationAdapter(manifest)

        with patch.object(adapter, "_dump_dom", return_value=dom_html):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 2)
        first = candidates[0]
        self.assertEqual(first.source_id, "rocketpunch")
        self.assertEqual(first.source_url, "https://www.rocketpunch.com/en/jobs?selectedJobId=158927")
        self.assertEqual(first.source_posting_id, "158927")
        self.assertEqual(first.title, "DataOps Engineer")
        self.assertEqual(first.company, "페이데이터")
        self.assertIn("Python", first.raw_jd["required_qualifications"])
        self.assertIn("data pipeline", first.raw_jd["responsibilities"][0])
        self.assertIn("Seoul", first.location)

    def test_rocketpunch_browser_automation_enriches_selected_job_detail(self) -> None:
        listing_url = "https://www.rocketpunch.com/en/jobs"
        manifest = SourceManifest(
            source_id="rocketpunch",
            enabled=True,
            access_mode="browser_automation",
            auth_required=False,
            tos_review_status="unknown",
            domains=["www.rocketpunch.com"],
            rate_limit="browser automation with user-directed notice override",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "search_urls": [listing_url],
                "candidate_include_keywords": ["python"],
                "policy_override_mode": "user_directed_ignore",
                "policy_override_reason": "User directed RocketPunch browser automation despite source notice.",
                "policy_override_acknowledges_source_notice": True,
                "delay_seconds": 0,
            },
        )
        listing_dom = """
        <div data-index="0">
          <a href="/en/jobs/158927?list=true">
            <p class="textStyle_Body.BodyS c_foregrounds.neutral.secondary">Playtica</p>
            <p class="textStyle_Body.BodyM_Bold c_foregrounds.neutral.primary">IT Specialist (BackEnd)</p>
            <p>Python, Java, Node.js</p>
          </a>
        </div>
        """
        detail_dom = (ROOT / "fixtures" / "chrome_captures" / "rocketpunch_selected_job_158927.html").read_text(
            encoding="utf-8"
        )
        adapter = RocketPunchBrowserAutomationAdapter(manifest)

        with patch.object(adapter, "_dump_dom", side_effect=[listing_dom, detail_dom]):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        candidate = candidates[0]
        self.assertEqual(candidate.source_url, "https://www.rocketpunch.com/en/jobs?selectedJobId=158927")
        self.assertEqual(candidate.deadline_raw, "2026-07-31")
        self.assertIn("unmanned/self-service", candidate.raw_jd["responsibilities"][0])
        self.assertIn("Django", candidate.raw_jd["required_qualifications"][0])
        self.assertIn("payment systems", candidate.raw_jd["preferred_qualifications"][0])
    def test_saramin_adapter_collects_from_official_api_payload(self) -> None:
        manifest = SourceManifest(
            source_id="saramin",
            enabled=True,
            access_mode="api",
            auth_required=False,
            tos_review_status="pass",
            domains=["oapi.saramin.co.kr"],
            rate_limit="official API quota",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "api_url": "https://oapi.saramin.co.kr/job-search",
                "access_key": "test-key",
                "approved_api_access": True,
                "require_robots": False,
                "candidate_include_keywords": ["python"],
                "delay_seconds": 0,
            },
        )
        adapter = SaraminAdapter(manifest)
        payload = {
            "jobs": {
                "job": [
                    {
                        "id": "saramin-1",
                        "url": "https://www.saramin.co.kr/zf_user/jobs/relay/view?rec_idx=1",
                        "expiration-date": "2026-07-31T23:59:59+09:00",
                        "company": {"detail": {"name": "Saramin AI"}},
                        "position": {
                            "title": "Python AI Engineer",
                            "location": {"name": "서울"},
                            "job-code": {"name": "AI·데이터"},
                            "experience-level": {"name": "경력 2년 이상"},
                        },
                    }
                ]
            }
        }

        with patch.object(
            adapter,
            "_get_fetch",
            return_value=HttpResponse("https://oapi.saramin.co.kr/job-search", json.dumps(payload)),
        ):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertEqual(candidates[0].source_id, "saramin")
        self.assertEqual(candidates[0].title, "Python AI Engineer")
        self.assertEqual(candidates[0].company, "Saramin AI")
        self.assertEqual(candidates[0].deadline_raw, "2026-07-31")
        self.assertEqual(candidates[0].raw_jd["experience_tags"], ["경력 2년 이상"])

    def test_saramin_adapter_collects_public_detail_body_without_api(self) -> None:
        detail_url = "https://www.saramin.co.kr/zf_user/jobs/relay/view-detail?rec_idx=54106686&rec_seq=0"
        manifest = SourceManifest(
            source_id="saramin",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.saramin.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "detail_urls": [detail_url],
                "candidate_include_keywords": ["python"],
                "require_robots": False,
                "explicit_automated_permission": True,
                "delay_seconds": 0,
            },
        )
        adapter = SaraminAdapter(manifest)
        detail_html = """
        <html><body>
        AI 엔지니어 채용
        사용 기술 Python FastAPI Redis AWS Docker
        주요업무
        LLM 에이전트 및 RAG 시스템 설계 및 개발
        자격요건
        신입 또는 관련 경력 1년 이상 Python 비동기 API 설계 경험
        우대사항
        Redis, Elasticsearch 운영 경험
        마감일 및 근무지
        마감일 : 2026년 07월 07일
        근무지
        - 서울 강남구 선릉로93길 40
        </body></html>
        """

        with patch.object(adapter, "_fetch", return_value=HttpResponse(detail_url, detail_html)):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertEqual(candidates[0].source_id, "saramin")
        self.assertEqual(candidates[0].source_posting_id, "54106686")
        self.assertEqual(candidates[0].title, "AI 엔지니어 채용")
        self.assertIn("LLM 에이전트", candidates[0].raw_jd["responsibilities"][0])
        self.assertIn("Python", candidates[0].raw_jd["required_qualifications"][0])
        self.assertIn("FastAPI", candidates[0].raw_jd["required_qualifications"])
        self.assertIn("Elasticsearch", candidates[0].raw_jd["preferred_qualifications"][0])
        self.assertIn("2026년 07월 07일", candidates[0].deadline_raw)
        self.assertIn("서울 강남구", candidates[0].location)

    def test_saramin_adapter_discovers_public_relay_detail_urls_from_listing(self) -> None:
        search_url = "https://www.saramin.co.kr/zf_user/search/recruit?searchword=python"
        detail_url = "https://www.saramin.co.kr/zf_user/jobs/relay/view-detail?rec_idx=54106686&rec_seq=0"
        manifest = SourceManifest(
            source_id="saramin",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.saramin.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "search_urls": [search_url],
                "link_include_keywords": ["python"],
                "candidate_include_keywords": ["python"],
                "require_robots": False,
                "explicit_automated_permission": True,
                "delay_seconds": 0,
            },
        )
        list_html = """
        <a href="/zf_user/jobs/relay/view?rec_idx=54106686&rec_seq=0">Python AI 엔지니어</a>
        <a href="/zf_user/jobs/relay/view?rec_idx=11111111&rec_seq=0">주방 보조</a>
        """
        detail_html = """
        <html><body>
        Python AI 엔지니어
        주요업무
        Python 기반 추천 시스템 개발
        자격요건
        Python API 개발 경험
        우대사항
        LLM 서비스 경험
        마감일 및 근무지
        근무지
        - 서울 강남구
        </body></html>
        """
        adapter = SaraminAdapter(manifest)

        def fake_fetch(url: str) -> HttpResponse:
            if url == search_url:
                return HttpResponse(url, list_html)
            self.assertEqual(url, detail_url)
            return HttpResponse(url, detail_html)

        with patch.object(adapter, "_fetch", side_effect=fake_fetch):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertEqual(candidates[0].source_url, detail_url)
        self.assertEqual(candidates[0].source_posting_id, "54106686")
        self.assertIn("추천 시스템", candidates[0].raw_jd["responsibilities"][0])

    def test_wanted_adapter_collects_user_provided_manual_postings(self) -> None:
        manifest = SourceManifest(
            source_id="wanted",
            enabled=True,
            access_mode="api",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.wanted.co.kr"],
            rate_limit="manual export",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "manual_postings": [
                    {
                        "url": "https://www.wanted.co.kr/wd/123",
                        "title": "Wanted Python ML Engineer",
                        "company": "Wanted AI",
                        "location": "Seoul",
                        "deadline": "2026-08-01",
                        "skills": ["Python", "LLM"],
                        "requirements": "Build machine learning products",
                        "experience": "경력 1년 이상",
                    }
                ],
                "candidate_include_keywords": ["python"],
                "delay_seconds": 0,
            },
        )
        candidates = WantedAdapter(manifest).collect()

        self.assertEqual(len(candidates), 1)
        self.assertEqual(candidates[0].source_id, "wanted")
        self.assertEqual(candidates[0].title, "Wanted Python ML Engineer")
        self.assertIn("LLM", candidates[0].raw_jd["required_qualifications"])

    def test_wanted_adapter_collects_public_detail_body_without_manual_payload(self) -> None:
        detail_url = "https://www.wanted.co.kr/wd/371139"
        manifest = SourceManifest(
            source_id="wanted",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.wanted.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "detail_urls": [detail_url],
                "candidate_include_keywords": ["python"],
                "require_robots": False,
                "explicit_automated_permission": True,
                "delay_seconds": 0,
            },
        )
        adapter = WantedAdapter(manifest)
        detail_html = """
        <html><head><meta property="og:title" content="[엑스보스] Python 워크플로우 자동화 개발자 채용 공고 | 원티드"></head><body>
        엑스보스∙서울 강서구∙경력 1-5년
        Python 워크플로우 자동화 개발자
        포지션 상세
        주요업무
        Python 워크플로우 설계와 API GUI 하이브리드 자동화 개발
        자격요건
        Python 개발 2-5년 또는 Selenium Playwright Requests asyncio 경험
        기술 스택 • 툴
        Git MySQL JSON Python Excel PostgreSQL SQLite
        마감일
        2026.07.25
        근무지역
        서울 강서구 마곡동 799-6
        본 채용정보는 원티드랩의 동의없이 무단전재
        </body></html>
        """

        with patch.object(adapter, "_fetch", return_value=HttpResponse(detail_url, detail_html)):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertEqual(candidates[0].source_id, "wanted")
        self.assertEqual(candidates[0].source_posting_id, "371139")
        self.assertEqual(candidates[0].title, "Python 워크플로우 자동화 개발자")
        self.assertEqual(candidates[0].company, "엑스보스")
        self.assertIn("Python 워크플로우", candidates[0].raw_jd["responsibilities"][0])
        self.assertIn("Selenium", candidates[0].raw_jd["required_qualifications"][0])
        self.assertIn("PostgreSQL", candidates[0].raw_jd["required_qualifications"])
        self.assertEqual(candidates[0].deadline_raw, "2026.07.25")
        self.assertIn("서울 강서구", candidates[0].location)

    def test_wanted_adapter_discovers_public_wd_urls_from_listing(self) -> None:
        search_url = "https://www.wanted.co.kr/search?query=python&tab=position"
        detail_url = "https://www.wanted.co.kr/wd/371139"
        manifest = SourceManifest(
            source_id="wanted",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.wanted.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "search_urls": [search_url],
                "link_include_keywords": ["python"],
                "candidate_include_keywords": ["python"],
                "require_robots": False,
                "explicit_automated_permission": True,
                "delay_seconds": 0,
            },
        )
        list_html = """
        <a href="/wd/371139">Python 워크플로우 자동화 개발자</a>
        <a href="/company/123">회사 소개</a>
        """
        detail_html = """
        <html><head><meta property="og:title" content="[엑스보스] Python 자동화 개발자 채용 공고 | 원티드"></head><body>
        엑스보스∙서울 강서구∙경력 1-5년
        Python 자동화 개발자
        주요업무
        Python 데이터 수집 자동화 개발
        자격요건
        Python Requests Playwright 경험
        마감일
        2026.07.25
        근무지역
        서울 강서구
        </body></html>
        """
        adapter = WantedAdapter(manifest)

        def fake_fetch(url: str) -> HttpResponse:
            if url == search_url:
                return HttpResponse(url, list_html)
            self.assertEqual(url, detail_url)
            return HttpResponse(url, detail_html)

        with patch.object(adapter, "_fetch", side_effect=fake_fetch):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertEqual(candidates[0].source_url, detail_url)
        self.assertEqual(candidates[0].source_posting_id, "371139")
        self.assertIn("데이터 수집", candidates[0].raw_jd["responsibilities"][0])
    def test_linkedin_adapter_collects_approved_partner_payload_path(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            payload_path = Path(tmp) / "linkedin_jobs.json"
            payload_path.write_text(
                json.dumps(
                    {
                        "jobs": [
                            {
                                "job_id": "li-1",
                                "job_url": "https://www.linkedin.com/jobs/view/1",
                                "job_title": "LinkedIn Data Engineer",
                                "company_name": "LinkedIn Partner Co",
                                "location": "Remote",
                                "deadline": "2026-08-15",
                                "skills": ["Python", "SQL", "data pipeline"],
                                "experience": "경력무관",
                            }
                        ]
                    }
                ),
                encoding="utf-8",
            )
            manifest = SourceManifest(
                source_id="linkedin",
                enabled=True,
                access_mode="api",
                auth_required=True,
                tos_review_status="pass",
                domains=["www.linkedin.com"],
                rate_limit="approved partner/API access only",
                failure_mode="skip_source",
                allowed_persisted_fields=[],
                options={
                    "approved_partner_access": True,
                    "approved_authenticated_flow": True,
                    "partner_payload_path": str(payload_path),
                    "candidate_include_keywords": ["python"],
                    "delay_seconds": 0,
                },
            )
            candidates = LinkedInAdapter(manifest).collect()

        self.assertEqual(len(candidates), 1)
        self.assertEqual(candidates[0].source_id, "linkedin")
        self.assertEqual(candidates[0].title, "LinkedIn Data Engineer")
        self.assertIn("data pipeline", candidates[0].raw_jd["required_qualifications"])

    def test_linkedin_partner_payload_requires_approved_access(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            payload_path = Path(tmp) / "linkedin_jobs.json"
            payload_path.write_text(
                json.dumps(
                    [
                        {
                            "job_url": "https://www.linkedin.com/jobs/view/1",
                            "job_title": "LinkedIn Data Engineer",
                            "company_name": "LinkedIn Partner Co",
                            "skills": ["Python"],
                        }
                    ]
                ),
                encoding="utf-8",
            )
            manifest = SourceManifest(
                source_id="linkedin",
                enabled=True,
                access_mode="api",
                auth_required=True,
                tos_review_status="pass",
                domains=["www.linkedin.com"],
                rate_limit="approved partner/API access only",
                failure_mode="skip_source",
                allowed_persisted_fields=[],
                options={
                    "approved_partner_access": False,
                    "approved_authenticated_flow": False,
                    "partner_payload_path": str(payload_path),
                },
            )

            with self.assertRaises(SourceAccessError):
                LinkedInAdapter(manifest).collect()

    def test_jobkorea_adapter_discovers_ai_jobs_from_api_html(self) -> None:
        manifest = SourceManifest(
            source_id="jobkorea",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.jobkorea.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "api_url": "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                "link_include_keywords": ["python"],
                "max_pages": 10,
                "require_robots": False,
                "explicit_automated_permission": True,
            },
        )
        adapter = JobKoreaAdapter(manifest)
        payload = {
            "html": """
            <a href="/Recruit/GI_Read/49333079?sc=323">AI 엔지니어 Python</a>
            <a href="/Recruit/GI_Read/49410271?sc=323">주방 보조</a>
            """
        }

        with patch.object(
            adapter,
            "_post_fetch",
            return_value=HttpResponse(
                "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                json.dumps(payload),
            ),
        ):
            urls = adapter.discover_urls()

        self.assertEqual(urls, ["https://www.jobkorea.co.kr/Recruit/GI_Read/49333079?sc=323"])

    def test_jobkorea_adapter_collects_api_keywords_as_snapshot_terms(self) -> None:
        manifest = SourceManifest(
            source_id="jobkorea",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.jobkorea.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "api_url": "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                "candidate_include_keywords": ["python"],
                "max_pages": 10,
                "require_robots": False,
                "explicit_automated_permission": True,
                "fetch_detail_pages": False,
                "delay_seconds": 0,
            },
        )
        adapter = JobKoreaAdapter(manifest)
        payload = {
            "html": """
            <li class="recruit-item">
              <a href="/Recruit/GI_Read/49333079?sc=323" class="recruit-link"
                 data-cname="Example AI" data-applyclosedt="2026-07-08 23:00:00">
                <div class="recruit-title"><h3 class="title">AI 엔지니어</h3></div>
                <ul class="keywords">
                  <li class="item primary">Python</li>
                  <li class="item primary">SQL</li>
                  <li class="item">서울 중구</li>
                </ul>
              </a>
            </li>
            </div></div>
            """
        }

        with patch.object(
            adapter,
            "_post_fetch",
            return_value=HttpResponse(
                "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                json.dumps(payload),
            ),
        ):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertEqual(candidates[0].company, "Example AI")
        self.assertEqual(candidates[0].deadline_raw, "2026-07-08")
        self.assertEqual(candidates[0].location, "서울 중구")
        self.assertIn("Python", candidates[0].raw_jd["required_qualifications"])
        self.assertNotIn("서울 중구", candidates[0].raw_jd["required_qualifications"])
        self.assertEqual(candidates[0].raw_jd["experience_tags"], [])

    def test_jobkorea_adapter_enriches_api_cards_with_public_detail_body(self) -> None:
        manifest = SourceManifest(
            source_id="jobkorea",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.jobkorea.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "api_url": "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                "candidate_include_keywords": ["python"],
                "max_pages": 10,
                "require_robots": False,
                "explicit_automated_permission": True,
                "delay_seconds": 0,
            },
        )
        adapter = JobKoreaAdapter(manifest)
        payload = {
            "html": """
            <li class="recruit-item">
              <a href="/Recruit/GI_Read/49476607?sc=323" class="recruit-link"
                 data-cname="플라잎" data-applyclosedt="2026-07-28 23:00:00">
                <div class="recruit-title"><h3 class="title">Forward Deployed 엔지니어</h3></div>
                <ul class="keywords">
                  <li class="item primary">Python</li>
                  <li class="item">울산 남구</li>
                </ul>
              </a>
            </li>
            """
        }
        detail_html = """
        <html><body>
          <h2>이런 업무를 해요</h2>
          <p>Python AI 시스템 개발과 고객 자동화 작업 분석</p>
          <h2>이런 분들을 찾고 있어요</h2>
          <p>PyTorch 경험 및 문제 해결 능력</p>
          <h2>우대사항</h2>
          <p>Physical AI 프로젝트 경험</p>
          <p>근무지 주소 : 울산광역시 남구 옥현로 129 지도보기</p>
        </body></html>
        """

        with patch.object(
            adapter,
            "_post_fetch",
            return_value=HttpResponse(
                "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                json.dumps(payload),
            ),
        ), patch.object(
            adapter,
            "_fetch",
            return_value=HttpResponse("https://www.jobkorea.co.kr/Recruit/GI_Read/49476607", detail_html),
        ):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertIn("고객 자동화", candidates[0].raw_jd["responsibilities"][0])
        self.assertIn("PyTorch", candidates[0].raw_jd["required_qualifications"][0])
        self.assertIn("Physical AI", candidates[0].raw_jd["preferred_qualifications"][0])
        self.assertEqual(candidates[0].location, "울산광역시 남구 옥현로 129")

    def test_jobkorea_adapter_uses_json_ld_when_detail_sections_are_absent(self) -> None:
        manifest = SourceManifest(
            source_id="jobkorea",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.jobkorea.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "api_url": "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                "candidate_include_keywords": ["python"],
                "max_pages": 10,
                "require_robots": False,
                "explicit_automated_permission": True,
                "require_detail_body": True,
                "delay_seconds": 0,
            },
        )
        adapter = JobKoreaAdapter(manifest)
        payload = {
            "html": """
            <li class="recruit-item">
              <a href="/Recruit/GI_Read/49333079?sc=323" class="recruit-link"
                 data-cname="Example AI" data-applyclosedt="2026-07-08 23:00:00">
                <div class="recruit-title"><h3 class="title">AI 엔지니어</h3></div>
                <ul class="keywords">
                  <li class="item primary">Python</li>
                  <li class="item">서울 중구</li>
                </ul>
              </a>
            </li>
            """
        }
        detail_html = """
        <html><head><title>Example AI 채용</title></head><body>
          <script type="application/ld+json">
          {
            "@context": "https://schema.org",
            "@type": "JobPosting",
            "title": "AI 엔지니어",
            "description": "Python 기반 LLM 서비스 개발자를 채용합니다.",
            "validThrough": "2026-07-08T23:59",
            "hiringOrganization": {"@type": "Organization", "name": "Example AI"},
            "jobLocation": {
              "@type": "Place",
              "address": {"@type": "PostalAddress", "streetAddress": "서울 중구 세종대로"}
            },
            "identifier": {"@type": "PropertyValue", "name": "JobKorea", "value": "49333079"},
            "url": "https://www.jobkorea.co.kr/Recruit/GI_Read/49333079"
          }
          </script>
          <p>섹션 마커 없는 Next page shell</p>
        </body></html>
        """

        with patch.object(
            adapter,
            "_post_fetch",
            return_value=HttpResponse(
                "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                json.dumps(payload),
            ),
        ), patch.object(
            adapter,
            "_fetch",
            return_value=HttpResponse("https://www.jobkorea.co.kr/Recruit/GI_Read/49333079", detail_html),
        ):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertIn("Python 기반 LLM", candidates[0].raw_jd["responsibilities"][0])
        self.assertEqual(candidates[0].location, "서울 중구 세종대로")
        self.assertEqual(candidates[0].deadline_raw, "2026-07-08")
        self.assertEqual(adapter.errors, [])

    def test_jobkorea_json_ld_graph_detail_quality_sample(self) -> None:
        manifest = SourceManifest(
            source_id="jobkorea",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.jobkorea.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "api_url": "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                "candidate_include_keywords": ["python"],
                "max_pages": 10,
                "require_robots": False,
                "explicit_automated_permission": True,
                "require_detail_body": True,
                "delay_seconds": 0,
            },
        )
        adapter = JobKoreaAdapter(manifest)
        payload = {
            "html": """
            <li class="recruit-item">
              <a href="/Recruit/GI_Read/50000001?sc=323" class="recruit-link"
                 data-cname="Graph AI" data-applyclosedt="2026-08-01 23:00:00">
                <div class="recruit-title"><h3 class="title">Python ML 엔지니어</h3></div>
                <ul class="keywords"><li class="item primary">Python</li></ul>
              </a>
            </li>
            """
        }
        detail_html = """
        <html><body>
          <script type="application/ld+json">
          {
            "@context": "https://schema.org",
            "@graph": [
              {"@type": "BreadcrumbList", "name": "ignored"},
              {
                "@type": "JobPosting",
                "title": "Python ML 엔지니어",
                "description": "<p>Python 모델 서빙과 ML pipeline 개발</p>",
                "validThrough": "2026-08-01T23:59:59+09:00",
                "hiringOrganization": {"@type": "Organization", "name": "Graph AI", "description": "AI product team"},
                "jobLocation": {"@type": "Place", "address": {"streetAddress": "서울 강남구 테헤란로 1"}}
              }
            ]
          }
          </script>
          <p>섹션 marker 없는 상세 shell</p>
        </body></html>
        """

        with patch.object(
            adapter,
            "_post_fetch",
            return_value=HttpResponse(
                "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                json.dumps(payload),
            ),
        ), patch.object(
            adapter,
            "_fetch",
            return_value=HttpResponse("https://www.jobkorea.co.kr/Recruit/GI_Read/50000001", detail_html),
        ):
            candidates = adapter.collect()

        self.assertEqual(len(candidates), 1)
        self.assertIn("ML pipeline", candidates[0].raw_jd["responsibilities"][0])
        self.assertIn("AI product team", candidates[0].raw_jd["company_info"])
        self.assertEqual(candidates[0].location, "서울 강남구 테헤란로 1")
        self.assertEqual(candidates[0].deadline_raw, "2026-08-01")

    def test_jobkorea_required_detail_body_drops_card_only_candidates(self) -> None:
        manifest = SourceManifest(
            source_id="jobkorea",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.jobkorea.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "api_url": "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                "candidate_include_keywords": ["python"],
                "max_pages": 10,
                "require_robots": False,
                "explicit_automated_permission": True,
                "require_detail_body": True,
                "delay_seconds": 0,
            },
        )
        adapter = JobKoreaAdapter(manifest)
        payload = {
            "html": """
            <li class="recruit-item">
              <a href="/Recruit/GI_Read/49476607?sc=323" class="recruit-link"
                 data-cname="플라잎" data-applyclosedt="2026-07-28 23:00:00">
                <div class="recruit-title"><h3 class="title">Python 엔지니어</h3></div>
              </a>
            </li>
            """
        }

        with patch.object(
            adapter,
            "_post_fetch",
            return_value=HttpResponse(
                "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                json.dumps(payload),
            ),
        ), patch.object(
            adapter,
            "_fetch",
            return_value=HttpResponse("https://www.jobkorea.co.kr/Recruit/GI_Read/49476607", "<html><body>카드만 있음</body></html>"),
        ):
            candidates = adapter.collect()

        self.assertEqual(candidates, [])
        self.assertTrue(any("detail body sections not found" in error for error in adapter.errors))

    def test_jobkorea_required_detail_body_rejects_empty_listing_fallback(self) -> None:
        manifest = SourceManifest(
            source_id="jobkorea",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.jobkorea.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "api_url": "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                "search_urls": ["https://www.jobkorea.co.kr/recruit/ai-jobs"],
                "candidate_include_keywords": ["python"],
                "max_pages": 10,
                "require_robots": False,
                "explicit_automated_permission": True,
                "require_detail_body": True,
                "delay_seconds": 0,
            },
        )
        adapter = JobKoreaAdapter(manifest)

        with patch.object(
            adapter,
            "_post_fetch",
            return_value=HttpResponse(
                "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                json.dumps({"html": ""}),
            ),
        ), patch.object(adapter, "_fetch") as fetch:
            candidates = adapter.collect()

        fetch.assert_not_called()
        self.assertEqual(candidates, [])
        self.assertEqual(adapter.errors, ["JobKorea listing API HTML was empty"])

    def test_jobkorea_adapter_preserves_experience_tags(self) -> None:
        manifest = SourceManifest(
            source_id="jobkorea",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.jobkorea.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "api_url": "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                "candidate_include_keywords": ["python"],
                "max_pages": 10,
                "require_robots": False,
                "explicit_automated_permission": True,
                "fetch_detail_pages": False,
                "delay_seconds": 0,
            },
        )
        adapter = JobKoreaAdapter(manifest)
        payload = {
            "html": """
            <li class="recruit-item">
              <a href="/Recruit/GI_Read/49333079?sc=323" class="recruit-link"
                 data-cname="Example AI" data-applyclosedt="2026-07-08 23:00:00">
                <div class="recruit-title"><h3 class="title">AI 엔지니어</h3></div>
                <ul class="keywords">
                  <li class="item primary">Python</li>
                  <li class="item">경력1년↑</li>
                  <li class="item">서울 중구</li>
                </ul>
              </a>
            </li>
            </div></div>
            """
        }

        with patch.object(
            adapter,
            "_post_fetch",
            return_value=HttpResponse(
                "https://www.jobkorea.co.kr/recruit/ai-jobs/GetRecruitList",
                json.dumps(payload),
            ),
        ):
            candidates = adapter.collect()

        self.assertEqual(candidates[0].raw_jd["experience_tags"], ["경력1년↑"])
        self.assertNotIn("경력1년↑", candidates[0].raw_jd["required_qualifications"])

    def test_live_run_allows_two_years_and_filters_above_profile_limit(self) -> None:
        raw = json.loads(CONFIG.read_text(encoding="utf-8"))
        raw["top_n"] = 5
        raw["profile"]["max_experience_years"] = 2
        raw["sources"] = [
            {
                "source_id": "fixture",
                "enabled": True,
                "access_mode": "manual",
                "auth_required": False,
                "tos_review_status": "not_required",
                "domains": [],
                "rate_limit": "none",
                "failure_mode": "skip_source",
                "allowed_persisted_fields": [],
            }
        ]
        postings = [
            {
                "source_id": "fixture",
                "source_url": "https://jobs.example.test/new-grad",
                "source_posting_id": "new-grad",
                "title": "New Grad AI Engineer",
                "company": "Example AI",
                "location": "Seoul",
                "deadline": "2026-07-10",
                "raw_jd": {
                    "required_qualifications": ["Python", "machine learning"],
                    "preferred_qualifications": ["PyTorch"],
                    "responsibilities": ["Build ML systems"],
                    "company_info": ["AI team"],
                    "experience_tags": ["경력무관"],
                },
            },
            {
                "source_id": "fixture",
                "source_url": "https://jobs.example.test/one-year",
                "source_posting_id": "one-year",
                "title": "AI Engineer 1 Year",
                "company": "Example AI",
                "location": "Seoul",
                "deadline": "2026-07-10",
                "raw_jd": {
                    "required_qualifications": ["Python", "machine learning"],
                    "preferred_qualifications": ["PyTorch"],
                    "responsibilities": ["Build ML systems"],
                    "company_info": ["AI team"],
                    "experience_tags": ["경력1년↑"],
                },
            },
            {
                "source_id": "fixture",
                "source_url": "https://jobs.example.test/three-year",
                "source_posting_id": "three-year",
                "title": "AI Engineer 3 Years",
                "company": "Example AI",
                "location": "Seoul",
                "deadline": "2026-07-10",
                "raw_jd": {
                    "required_qualifications": ["Python", "machine learning"],
                    "preferred_qualifications": ["PyTorch"],
                    "responsibilities": ["Build ML systems"],
                    "company_info": ["AI team"],
                    "experience_tags": ["경력3년↑"],
                },
            },
        ]
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            config_path = tmp_path / "config.json"
            fixture_path = tmp_path / "postings.json"
            raw["fixture_path"] = str(fixture_path)
            raw["output_dir"] = str(tmp_path / "reports")
            config_path.write_text(json.dumps(raw), encoding="utf-8")
            fixture_path.write_text(json.dumps(postings), encoding="utf-8")
            config = load_config(config_path, allow_real_sources=True)
            summary, report, ranked = run_live_run(config, date(2026, 6, 30))

        self.assertEqual(summary.experience_excluded, 1)
        self.assertEqual(len(ranked), 2)
        self.assertIn("New Grad AI Engineer", report)
        self.assertIn("AI Engineer 1 Year", report)
        self.assertNotIn("AI Engineer 3 Years", report)

    def test_disabling_robots_requires_explicit_permission(self) -> None:
        manifest = SourceManifest(
            source_id="wanted",
            enabled=True,
            access_mode="public_page",
            auth_required=False,
            tos_review_status="pass",
            domains=["www.wanted.co.kr"],
            rate_limit="1 request / second",
            failure_mode="skip_source",
            allowed_persisted_fields=[],
            options={
                "start_urls": ["https://www.wanted.co.kr/wd/1"],
                "require_robots": False,
            },
        )
        adapter = PublicJobsHttpAdapter(manifest)

        with self.assertRaises(SourceAccessError):
            adapter.collect()

class SourceRegistryTests(unittest.TestCase):
    def _live_config(self) -> dict:
        return json.loads((ROOT / "config" / "live_sources.sample.json").read_text(encoding="utf-8"))

    def _write_temp_config(self, raw: dict, tmp_path: Path) -> Path:
        config_path = tmp_path / "live_sources.json"
        config_path.write_text(json.dumps(raw), encoding="utf-8")
        return config_path

    def test_source_registry_loads_expected_statuses_and_lanes(self) -> None:
        config = load_config(ROOT / "config" / "live_sources.sample.json", allow_real_sources=True)
        by_id = {source.source_id: source for source in config.sources}

        self.assertEqual(
            set(by_id),
            {
                "company_careers",
                "saramin",
                "jobkorea",
                "wanted",
                "jumpit",
                "rallit",
                "rocketpunch",
                "linkedin",
                "naver_careers",
                "kakao_careers",
                "line_careers",
                "coupang_careers",
            },
        )
        self.assertEqual(by_id["jumpit"].target_lane, "public_http")
        self.assertEqual(by_id["jumpit"].target_status, "enabled")
        self.assertEqual(by_id["rallit"].target_lane, "public_http")
        self.assertEqual(by_id["rallit"].target_status, "enabled")
        self.assertEqual(by_id["jobkorea"].target_status, "enabled")
        self.assertEqual(by_id["jobkorea"].target_lane, "public_http")
        self.assertEqual(by_id["saramin"].target_status, "enabled")
        self.assertEqual(by_id["saramin"].target_lane, "public_http")
        self.assertEqual(by_id["wanted"].target_status, "enabled")
        self.assertEqual(by_id["wanted"].target_lane, "public_http")
        self.assertEqual(by_id["rocketpunch"].target_status, "enabled")
        self.assertEqual(by_id["rocketpunch"].target_lane, "browser_automation")
        self.assertEqual(by_id["linkedin"].target_status, "excluded")
        self.assertEqual(by_id["naver_careers"].target_status, "deferred")
        self.assertIsNone(by_id["naver_careers"].target_lane)
        self.assertEqual(by_id["kakao_careers"].target_status, "deferred")
        self.assertIsNone(by_id["kakao_careers"].target_lane)
        self.assertEqual(by_id["line_careers"].target_status, "deferred")
        self.assertIsNone(by_id["line_careers"].target_lane)

    def test_source_registry_rejects_empty_target_lane(self) -> None:
        raw = self._live_config()
        raw["sources"][0]["target_lane"] = ""
        with tempfile.TemporaryDirectory() as tmp:
            with self.assertRaisesRegex(ConfigError, "target_lane cannot be empty string"):
                load_config(self._write_temp_config(raw, Path(tmp)), allow_real_sources=True)

    def test_source_registry_rejects_non_target_target_enablement(self) -> None:
        blocked_options = [
            ("saramin", {"access_mode": "api", "options": {"approved_api_access": True}}),
            ("wanted", {"options": {"manual_postings": [{"title": "Manual"}]}}),
            ("wanted", {"options": {"manual_export_path": "wanted.csv"}}),
            ("wanted", {"options": {"user_operated_chrome_extension": True}}),
            ("wanted", {"options": {"user_operated_browser_use": True}}),
            ("saramin", {"options": {"ocr_required": True}}),
            ("saramin", {"options": {"manual_review_flags": ["본문 OCR 필요"]}}),
            ("linkedin", {"options": {"approved_partner_access": True, "partner_payload_path": "jobs.json"}}),
        ]
        for source_id, override in blocked_options:
            raw = self._live_config()
            for source in raw["sources"]:
                if source["source_id"] == source_id:
                    source.update(
                        {
                            "enabled": True,
                            "target_status": "enabled",
                            "target_lane": "public_http",
                            "automation_level": "no_human",
                            "tos_review_status": "pass",
                            "adapter_code_path": "src/recruit_crawler/sources/platforms.py::Adapter",
                            "test_refs": ["tests/test_dry_run.py::example"],
                            "docs_refs": ["docs/source_collection_matrix.md"],
                        }
                    )
                    source.update({key: value for key, value in override.items() if key != "options"})
                    source["options"].update(override.get("options", {}))
            with tempfile.TemporaryDirectory() as tmp:
                with self.assertRaises(ConfigError, msg=source_id):
                    load_config(self._write_temp_config(raw, Path(tmp)), allow_real_sources=True)

    def test_registry_allows_user_directed_policy_override_for_browser_automation(self) -> None:
        raw = self._live_config()
        for source in raw["sources"]:
            if source["source_id"] == "rocketpunch":
                source.update(
                    {
                        "enabled": True,
                        "access_mode": "browser_automation",
                        "target_status": "enabled",
                        "target_lane": "browser_automation",
                        "automation_level": "no_human",
                        "tos_review_status": "unknown",
                        "adapter_code_path": "src/recruit_crawler/sources/platforms.py::RocketPunchBrowserAutomationAdapter",
                        "test_refs": ["tests/test_dry_run.py::test_rocketpunch_browser_automation_parses_listing_cards_without_detail_links"],
                        "docs_refs": ["docs/source_search_logic.md"],
                    }
                )
                source["options"].update(
                    {
                        "policy_override_mode": "user_directed_ignore",
                        "policy_override_reason": "User directed RocketPunch no-human browser automation despite source notice.",
                        "policy_override_acknowledges_source_notice": True,
                    }
                )

        with tempfile.TemporaryDirectory() as tmp:
            config = load_config(self._write_temp_config(raw, Path(tmp)), allow_real_sources=True)

        rocketpunch = next(source for source in config.sources if source.source_id == "rocketpunch")
        self.assertEqual(rocketpunch.target_lane, "browser_automation")
        self.assertEqual(rocketpunch.options["policy_override_mode"], "user_directed_ignore")

    def test_policy_override_requires_explicit_acknowledgement(self) -> None:
        raw = self._live_config()
        for source in raw["sources"]:
            if source["source_id"] == "rocketpunch":
                source.update(
                    {
                        "enabled": True,
                        "access_mode": "browser_automation",
                        "target_status": "enabled",
                        "target_lane": "browser_automation",
                        "automation_level": "no_human",
                        "tos_review_status": "unknown",
                        "adapter_code_path": "src/recruit_crawler/sources/platforms.py::RocketPunchBrowserAutomationAdapter",
                        "test_refs": ["tests/test_dry_run.py::test_rocketpunch_browser_automation_parses_listing_cards_without_detail_links"],
                        "docs_refs": ["docs/source_search_logic.md"],
                    }
                )
                source["options"].update(
                    {
                        "policy_override_mode": "user_directed_ignore",
                        "policy_override_reason": "User directed RocketPunch no-human browser automation despite source notice.",
                        "policy_override_acknowledges_source_notice": False,
                    }
                )

        with tempfile.TemporaryDirectory() as tmp:
            with self.assertRaisesRegex(ConfigError, "requires passed source review"):
                load_config(self._write_temp_config(raw, Path(tmp)), allow_real_sources=True)

    def test_enabled_registry_sources_require_code_tests_and_docs_refs(self) -> None:
        raw = self._live_config()
        for source in raw["sources"]:
            if source["source_id"] == "jumpit":
                source["test_refs"] = []
        with tempfile.TemporaryDirectory() as tmp:
            with self.assertRaisesRegex(ConfigError, "requires test_refs"):
                load_config(self._write_temp_config(raw, Path(tmp)), allow_real_sources=True)
    def test_registry_docs_refs_exist_for_every_source(self) -> None:
        config = load_config(ROOT / "config" / "live_sources.sample.json", allow_real_sources=True)

        for source in config.sources:
            for docs_ref in source.docs_refs:
                self.assertTrue((ROOT / docs_ref).exists(), f"{source.source_id}: {docs_ref}")

    def test_registry_test_refs_exist_for_every_enabled_source(self) -> None:
        config = load_config(ROOT / "config" / "live_sources.sample.json", allow_real_sources=True)
        test_file = ROOT / "tests" / "test_dry_run.py"
        test_text = test_file.read_text(encoding="utf-8")

        for source in config.sources:
            if source.target_status != "enabled":
                continue
            for test_ref in source.test_refs:
                test_path, _, test_name = test_ref.partition("::")
                self.assertTrue((ROOT / test_path).exists(), f"{source.source_id}: {test_ref}")
                self.assertIn(f"def {test_name}(", test_text, f"{source.source_id}: {test_ref}")


class SourceStatusCliTests(unittest.TestCase):
    def test_source_status_json_outputs_registry_without_network_or_adapter_construction(self) -> None:
        stdout = io.StringIO()
        with patch("recruit_crawler.sources.base.build_source_adapter") as build_adapter:
            with redirect_stdout(stdout):
                exit_code = cli_main(
                    [
                        "source-status",
                        "--config",
                        str(ROOT / "config" / "live_sources.sample.json"),
                        "--json",
                    ]
                )

        self.assertEqual(exit_code, 0)
        build_adapter.assert_not_called()
        payload = json.loads(stdout.getvalue())
        by_id = {source["source_id"]: source for source in payload["sources"]}
        self.assertEqual(
            set(by_id),
            {
                "company_careers",
                "saramin",
                "jobkorea",
                "wanted",
                "jumpit",
                "rallit",
                "rocketpunch",
                "linkedin",
                "naver_careers",
                "kakao_careers",
                "line_careers",
                "coupang_careers",
            },
        )
        self.assertEqual(payload["sources"][-1]["source_id"], "coupang_careers")
        for source in by_id.values():
            self.assertIn(source["target_lane"], {"public_http", "browser_automation", None})
        self.assertEqual(by_id["jumpit"]["target_lane"], "public_http")
        self.assertEqual(by_id["rallit"]["target_lane"], "public_http")
        self.assertEqual(by_id["jobkorea"]["target_lane"], "public_http")
        self.assertEqual(by_id["saramin"]["target_lane"], "public_http")
        self.assertEqual(by_id["wanted"]["target_lane"], "public_http")
        self.assertEqual(by_id["rocketpunch"]["target_lane"], "browser_automation")
        self.assertIsNone(by_id["naver_careers"]["target_lane"])
        self.assertIsNone(by_id["kakao_careers"]["target_lane"])
        self.assertEqual(by_id["rocketpunch"]["target_status"], "enabled")
        self.assertIsNone(by_id["linkedin"]["target_lane"])

class ChromeExtensionBoundaryTests(unittest.TestCase):
    def test_static_content_script_load_is_passive(self) -> None:
        content = (ROOT / "browser_extension" / "content.js").read_text(encoding="utf-8")

        self.assertIn('const CAPTURE_COMMAND = "recruit-capture:capture-visible-postings";', content)
        self.assertIn("registerCaptureCommandHandler();", content)
        self.assertIn("injectCaptureButton();", content)
        self.assertNotIn("return captureVisiblePostings();", content)
        self.assertLess(content.rfind("registerCaptureCommandHandler();"), content.rfind("injectCaptureButton();"))
        self.assertNotIn("window.addEventListener(\"message\"", content)
        self.assertNotIn("postMessage({", content)

    def test_popup_uses_explicit_capture_command(self) -> None:
        popup = (ROOT / "browser_extension" / "popup.js").read_text(encoding="utf-8")

        self.assertIn('"recruit-capture:capture-visible-postings"', popup)
        self.assertIn("chrome.tabs.sendMessage", popup)
        self.assertIn("frameId: 0", popup)
        self.assertNotIn('type: "recruit-capture:download"', popup)
        self.assertNotIn("const [{ result }]", popup)

    def test_capture_payload_includes_diagnostics_and_download_proof(self) -> None:
        content = (ROOT / "browser_extension" / "content.js").read_text(encoding="utf-8")
        background = (ROOT / "browser_extension" / "background.js").read_text(encoding="utf-8")

        self.assertIn('const EXTENSION_VERSION = "0.1.0";', content)
        self.assertIn("function withCaptureDiagnostics", content)
        self.assertIn("extension_version: EXTENSION_VERSION", content)
        self.assertIn("detail_length: requirements.length", content)
        self.assertIn("marker_hit: sourceDetailMarkerPattern().test(requirements)", content)
        self.assertIn('extraction_strategy: "linkedin_visible_detail_clickthrough"', content)
        self.assertIn("clickthrough:", content)
        self.assertIn("iframe_status: \"same_origin_dom_only\"", content)
        self.assertIn("filename", background)
        self.assertIn("sendResponse({ ok: true, ...download })", background)

class CaptureImportTests(unittest.TestCase):
    def _write_config(self, tmp_path: Path) -> Path:
        raw = json.loads(CONFIG.read_text(encoding="utf-8"))
        raw["output_dir"] = str(tmp_path / "reports")
        config_path = tmp_path / "config.json"
        config_path.write_text(json.dumps(raw), encoding="utf-8")
        return config_path

    def test_capture_import_maps_mixed_sources_and_generates_report(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            day_dir = tmp_path / "spool" / "2026-06-30"
            for source in ("linkedin", "saramin", "jobkorea"):
                (day_dir / source).mkdir(parents=True)
            captures = [
                {
                    "source_id": "linkedin",
                    "captured_at": "2026-06-30T04:00:00Z",
                    "postings": [
                        {
                            "source_id": "linkedin",
                            "source_url": "https://www.linkedin.com/jobs/view/4432928554/",
                            "source_posting_id": "4432928554",
                            "title": "LinkedIn Data Engineer",
                            "company": "LinkedIn Partner Co",
                            "location": "서울 서울",
                            "deadline": "",
                            "skills": ["Python", "SQL"],
                            "requirements": "Minimum Qualifications Python Programming Deep learning with PyTorch",
                            "captured_at": "2026-06-30T04:00:00Z",
                            "unexpected_private_note": "SHOULD_NOT_LEAK",
                        }
                    ],
                },
                {
                    "source_id": "saramin",
                    "captured_at": "2026-06-30T04:01:00Z",
                    "postings": [
                        {
                            "source_id": "saramin",
                            "source_url": "https://www.saramin.co.kr/zf_user/jobs/relay/view?rec_idx=54106686",
                            "source_posting_id": "54106686",
                            "title": "AI 엔지니어 채용",
                            "company": "레플리",
                            "location": "서울 강남구",
                            "deadline": "",
                            "skills": ["신입", "AI", "Python"],
                            "requirements": "주요업무 LLM 에이전트 및 RAG 시스템 설계 자격요건 Python FastAPI 마감일 : 2026년 07월 07일",
                            "captured_at": "2026-06-30T04:01:00Z",
                        }
                    ],
                },
                {
                    "source_id": "jobkorea",
                    "captured_at": "2026-06-30T04:02:00Z",
                    "postings": [
                        {
                            "source_id": "jobkorea",
                            "source_url": "https://www.jobkorea.co.kr/Recruit/GI_Read/49476607?sc=322",
                            "source_posting_id": "49476607",
                            "title": "[울산]Forward Deployed 엔지니어",
                            "company": "플라잎",
                            "location": "울산 남구 마감일 ~7/28(화",
                            "deadline": "",
                            "skills": ["AI", "Python"],
                            "requirements": "이런 업무를 해요 고객 자동화 작업 분석 이런 분들을 찾고 있어요 PyTorch Jax 근무지 주소 : 울산광역시 남구 옥현로 129 지도보기 이 기간동안 모집해요 ~ 2026.07.28(화)",
                            "captured_at": "2026-06-30T04:02:00Z",
                        },
                        {
                            "source_id": "jobkorea",
                            "source_url": "https://www.jobkorea.co.kr/Recruit/GI_Read/49476607?sc=322",
                            "source_posting_id": "49476607",
                            "title": "Duplicate",
                            "company": "플라잎",
                            "location": "울산 남구",
                            "skills": ["Python"],
                            "requirements": "duplicate",
                        },
                    ],
                },
                {"source_id": "saramin", "captured_at": "2026-06-30T04:03:00Z", "postings": []},
            ]
            for index, capture in enumerate(captures):
                source = capture["source_id"]
                (day_dir / source / f"capture-{index}.json").write_text(json.dumps(capture), encoding="utf-8")
            (day_dir / "jobkorea" / "invalid.json").write_text("{not-json", encoding="utf-8")

            selection = select_capture_files(tmp_path / "spool", run_date=date(2026, 6, 30))
            imported = import_capture_files(selection.files)
            config = load_config(self._write_config(tmp_path))
            summary, report, ranked = run_capture_import(
                config,
                selection.run_date,
                imported.candidates,
                imported.sources_attempted,
                imported.source_errors,
            )

        self.assertEqual(summary.candidates_collected, 3)
        self.assertEqual(summary.duplicates_removed, 0)
        self.assertEqual(summary.ranked_count, 3)
        self.assertEqual(imported.sources_attempted, ["jobkorea", "linkedin", "saramin"])
        self.assertTrue(any("invalid JSON" in error for error in summary.source_errors))
        self.assertTrue(any("empty postings" in error for error in summary.source_errors))
        self.assertTrue(any("duplicate posting jobkorea:49476607" in error for error in summary.source_errors))
        self.assertIn("LinkedIn Data Engineer", report)
        self.assertIn("AI 엔지니어 채용", report)
        self.assertIn("[울산]Forward Deployed 엔지니어", report)
        self.assertIn("울산광역시 남구 옥현로 129", report)
        self.assertIn("2026-07-28", report)
        self.assertNotIn("울산 남구 마감일", report)
        self.assertNotIn("SHOULD_NOT_LEAK", report)
        self.assertIn("PyTorch", report)
        self.assertIn("FastAPI", report)
        self.assertEqual({item.snapshot.source_id for item in ranked}, {"linkedin", "saramin", "jobkorea"})

    def test_capture_import_rejects_sensitive_posting_fields(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "capture.json"
            path.write_text(
                json.dumps(
                    {
                        "source_id": "linkedin",
                        "postings": [
                            {
                                "source_id": "linkedin",
                                "source_url": "https://www.linkedin.com/jobs/view/1",
                                "title": "Data Engineer",
                                "company": "Example",
                                "session_token": "secret",
                            }
                        ],
                    }
                ),
                encoding="utf-8",
            )

            imported = import_capture_files([path])

        self.assertEqual(imported.candidates, [])
        self.assertTrue(any("sensitive field" in error for error in imported.source_errors))

    def test_saramin_image_only_capture_is_marked_for_manual_ocr_review(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            path = tmp_path / "capture.json"
            path.write_text(
                json.dumps(
                    {
                        "source_id": "saramin",
                        "postings": [
                            {
                                "source_id": "saramin",
                                "source_url": "https://www.saramin.co.kr/zf_user/jobs/relay/view?rec_idx=1",
                                "source_posting_id": "1",
                                "title": "Image JD Engineer",
                                "company": "Image Co",
                                "location": "서울",
                                "skills": ["Python"],
                                "requirements": "",
                            }
                        ],
                    }
                ),
                encoding="utf-8",
            )
            imported = import_capture_files([path])
            raw = json.loads(CONFIG.read_text(encoding="utf-8"))
            raw["output_dir"] = str(tmp_path / "reports")
            config_path = tmp_path / "config.json"
            config_path.write_text(json.dumps(raw), encoding="utf-8")
            config = load_config(config_path)
            _summary, report, ranked = run_capture_import(
                config,
                date(2026, 6, 30),
                imported.candidates,
                imported.sources_attempted,
                imported.source_errors,
            )

        self.assertEqual(ranked[0].snapshot.manual_review_flags, ["본문 OCR 필요: 사람인 이미지형 JD 또는 DOM 텍스트 없음"])
        self.assertIn("본문 OCR 필요", report)
        self.assertIn("본문 이미지/OCR 필요 상태를 수동 검토했나요?", report)

    def test_capture_quality_gate_reports_privacy_and_import_categories(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            good = tmp_path / "good.json"
            bad = tmp_path / "bad.json"
            good.write_text(
                json.dumps(
                    {
                        "source_id": "saramin",
                        "captured_at": "2026-06-30T04:00:00Z",
                        "postings": [
                            {
                                "source_id": "saramin",
                                "source_url": "https://www.saramin.co.kr/zf_user/jobs/relay/view?rec_idx=2",
                                "source_posting_id": "2",
                                "title": "Public Contact Engineer",
                                "company": "Contact Co",
                                "location": "서울",
                                "skills": ["Python"],
                                "requirements": "자격요건 Python 문의 recruit@example.com",
                            },
                            {
                                "source_id": "saramin",
                                "source_url": "https://www.saramin.co.kr/zf_user/jobs/relay/view?rec_idx=3",
                                "source_posting_id": "3",
                                "title": "Image JD Engineer",
                                "company": "Image Co",
                                "location": "서울",
                                "skills": ["Python"],
                                "requirements": "",
                            },
                        ],
                    }
                ),
                encoding="utf-8",
            )
            bad.write_text("{not-json", encoding="utf-8")
            selection = select_capture_files(tmp_path, files=[good, bad], run_date=date(2026, 6, 30))
            imported = import_capture_files(selection.files)
            gate = build_capture_quality_gate(selection, imported)

        self.assertEqual(gate["status"], "fail")
        self.assertTrue(any(item["severity"] == "fail" and "invalid JSON" in item["message"] for item in gate["findings"]))
        self.assertTrue(any(item["category"] == "warning" and "public JD contact" in item["message"] for item in gate["privacy"]))
        self.assertEqual(gate["manual_review_items"][0]["flags"], ["본문 OCR 필요: 사람인 이미지형 JD 또는 DOM 텍스트 없음"])
        self.assertEqual(gate["source_mode_counts"], {"saramin": 2})

    def test_checked_in_chrome_capture_fixtures_cover_source_modes(self) -> None:
        fixture_paths = [
            ROOT / "fixtures" / "chrome_captures" / "linkedin_detail.json",
            ROOT / "fixtures" / "chrome_captures" / "saramin_image_only.json",
            ROOT / "fixtures" / "chrome_captures" / "jobkorea_detail.json",
        ]
        selection = select_capture_files(ROOT / "fixtures" / "chrome_captures", files=fixture_paths, run_date=date(2026, 6, 30))
        imported = import_capture_files(selection.files)
        gate = build_capture_quality_gate(selection, imported)

        self.assertEqual(gate["source_mode_counts"], {"jobkorea": 1, "linkedin": 1, "saramin": 1})
        self.assertEqual(gate["status"], "manual_review_required")
        self.assertTrue(gate["manual_review_items"])
        locations = {candidate.source_id: candidate.location for candidate in imported.candidates}
        self.assertEqual(locations["jobkorea"], "울산광역시 남구 옥현로 129")

class UserContextSchemaTests(unittest.TestCase):
    def test_config_profile_creates_user_context_contract(self) -> None:
        config = load_config(CONFIG)

        self.assertEqual(config.user_context.skills, config.profile.skills)
        self.assertEqual(config.user_context.explicit_deal_breakers, config.profile.exclusions)
        self.assertEqual(config.user_context.provenance["skills"], "config.profile.skills")

    def test_explicit_deal_breaker_excludes_without_raw_leakage(self) -> None:
        config = load_config(CONFIG)
        snapshot = JDSnapshot(
            source_id="fixture",
            source_url="https://jobs.example.test/unpaid",
            source_posting_id="unpaid",
            title="Unpaid Internship ML Engineer",
            company="Example",
            location="Seoul",
            deadline_raw=None,
            deadline=None,
            deadline_uncertain=False,
            required_qualifications=["Python", "machine learning"],
            preferred_qualifications=["LLM"],
            responsibilities=["Build models"],
            company_info=["AI team"],
        )

        assessment = score_snapshot(snapshot, config)

        self.assertEqual(assessment.verdict, "exclude")
        self.assertEqual(assessment.recommendation, "low_priority")
        self.assertEqual(assessment.deal_breaker_hits, ["unpaid internship"])
        self.assertNotIn("PRIVATE_PROFILE_CANARY", " ".join(assessment.risks))


class DocumentContextParserTests(unittest.TestCase):
    def test_plaintext_context_imports_user_context(self) -> None:
        context = parse_context_document(ROOT / "fixtures" / "user_context" / "context.md")

        self.assertIn("Python", context.skills)
        self.assertIn("Seoul", context.preferred_locations)
        self.assertEqual(context.max_experience_years, 2)

    def test_private_canary_document_fails_closed(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "private.md"
            path.write_text("Skills: Python\nPRIVATE_PROFILE_CANARY", encoding="utf-8")

            with self.assertRaises(UserContextImportError):
                parse_context_document(path)

    def test_docx_context_imports_user_context_when_dependency_available(self) -> None:
        try:
            from docx import Document
        except Exception:
            self.skipTest("python-docx is not installed")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "context.docx"
            doc = Document()
            doc.add_paragraph("Roles: ML Engineer")
            doc.add_paragraph("Skills: Python, SQL")
            doc.add_paragraph("Locations: Remote")
            doc.add_paragraph("Experience: 2 years")
            doc.save(str(path))

            context = parse_context_document(path)

        self.assertIn("SQL", context.skills)
        self.assertIn("Remote", context.preferred_locations)

    def test_pdf_context_imports_user_context_when_dependency_available(self) -> None:
        try:
            from pypdf import PdfWriter
        except Exception:
            self.skipTest("pypdf is not installed")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "empty.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=72, height=72)
            with path.open("wb") as fh:
                writer.write(fh)

            with self.assertRaises(UserContextImportError):
                parse_context_document(path)


class SupplementalInterviewTests(unittest.TestCase):
    def test_missing_context_generates_questions_and_answers_merge(self) -> None:
        context = UserContext(desired_roles=[], skills=[], preferred_locations=[], max_experience_years=0)

        questions = supplemental_questions(context)
        merged = merge_supplemental_answers(
            context,
            {
                "desired_roles": "ML Engineer",
                "skills": "Python, SQL",
                "preferred_locations": "Seoul",
                "max_experience_years": "2",
            },
        )

        self.assertGreaterEqual(len(questions), 4)
        self.assertEqual(merged.skills, ["Python", "SQL"])
        self.assertEqual(merged.provenance["skills"], "supplemental_interview")


class RelevanceCaseEvaluationTests(unittest.TestCase):
    def _snapshot(self, case_id: int, *, title: str, required: list[str], location: str = "Seoul") -> JDSnapshot:
        return JDSnapshot(
            source_id="fixture",
            source_url=f"https://jobs.example.test/{case_id}",
            source_posting_id=str(case_id),
            title=title,
            company="Example",
            location=location,
            deadline_raw=None,
            deadline=None,
            deadline_uncertain=False,
            required_qualifications=required,
            preferred_qualifications=["LLM"],
            responsibilities=["Build Python data products"],
            company_info=["AI team"],
        )

    def test_thirty_seed_relevance_cases_evaluate_deterministically(self) -> None:
        config = load_config(CONFIG)
        cases = []
        for index in range(30):
            if index % 3 == 0:
                context = config.user_context
                expected = "include"
                snapshot = self._snapshot(index, title="ML Engineer", required=["Python", "machine learning"])
                movement = "up"
            elif index % 3 == 1:
                context = UserContext(desired_roles=["ML Engineer"], skills=[], preferred_locations=["Seoul"], max_experience_years=2)
                expected = "hold"
                snapshot = self._snapshot(index, title="ML Engineer", required=["Rust"])
                movement = "same"
            else:
                context = UserContext(
                    desired_roles=["ML Engineer"],
                    skills=["Python"],
                    preferred_locations=["Seoul"],
                    max_experience_years=2,
                    explicit_deal_breakers=["unpaid internship"],
                )
                expected = "exclude"
                snapshot = self._snapshot(index, title="Unpaid Internship ML Engineer", required=["Python"])
                movement = "down"
            cases.append(
                RelevanceCase(
                    case_id=f"case-{index:02d}",
                    user_context=context,
                    snapshot=snapshot,
                    expected_verdict=expected,
                    expected_movement=movement,
                )
            )

        failures = evaluate_relevance_cases(cases, config)

        self.assertEqual(len(cases), 30)
        self.assertEqual(failures, [])


class SourceCandidateSelectionTests(unittest.TestCase):
    def test_company_careers_gate_keeps_zero_candidate_targets_deferred(self) -> None:
        config = load_config(ROOT / "config" / "live_sources.sample.json", allow_real_sources=True)
        selected = [source for source in config.sources if source.v1_role == "company_careers_selected"]
        fallbacks = [source for source in config.sources if source.v1_role == "company_careers_fallback"]

        self.assertEqual({source.source_id for source in selected}, {"naver_careers", "kakao_careers"})
        self.assertEqual({source.source_id for source in fallbacks}, {"line_careers", "coupang_careers"})
        for source in selected:
            self.assertFalse(source.enabled)
            self.assertIsNone(source.target_lane)
            self.assertEqual(source.target_status, "deferred")
            self.assertEqual(source.access_mode, "public_page")
            self.assertTrue(source.blockers)
            self.assertFalse(source.auth_required)


class BrowserEvidenceCliTests(unittest.TestCase):
    def test_browser_evidence_fixture_writes_allowed_fields_without_dom_leakage(self) -> None:
        output = io.StringIO()
        with tempfile.TemporaryDirectory() as tmp, redirect_stdout(output):
            exit_code = cli_main(
                [
                    "browser-evidence",
                    "--config",
                    str(ROOT / "config" / "live_sources.sample.json"),
                    "--source-id",
                    "rocketpunch",
                    "--fixture-html",
                    str(ROOT / "fixtures" / "browser_evidence" / "rocketpunch_listing.html"),
                    "--output",
                    str(Path(tmp) / "rocketpunch_fixture.json"),
                ]
            )
            transcript = json.loads((Path(tmp) / "rocketpunch_fixture.json").read_text(encoding="utf-8"))

        self.assertEqual(exit_code, 0)
        self.assertEqual(transcript["schema_version"], 1)
        self.assertEqual(transcript["command_mode"], "fixture")
        self.assertIn("dom_sha256", transcript)
        self.assertNotIn("dom", transcript)
        self.assertEqual(transcript["privacy_findings"], [])
        self.assertTrue(transcript["filterability"]["stable_posting_url"])

    def test_browser_evidence_private_target_fails(self) -> None:
        config = load_config(ROOT / "config" / "live_sources.sample.json", allow_real_sources=True)
        rocketpunch = next(source for source in config.sources if source.source_id == "rocketpunch")

        transcript = build_browser_evidence(rocketpunch, target_url="https://www.rocketpunch.com/private?session=secret")

        self.assertEqual(transcript["exit_code"], 1)
        self.assertTrue(transcript["errors"])

    def test_browser_evidence_redacts_private_markers_case_insensitively(self) -> None:
        self.assertEqual(
            _redact("https://jobs.example.test/apply?Session=abc&ACCESS_TOKEN=def"),
            "https://jobs.example.test/apply?[REDACTED]abc&[REDACTED]def",
        )


if __name__ == "__main__":
    unittest.main()
