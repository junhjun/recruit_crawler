from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from recruit_crawler.config import load_config
from recruit_crawler.scorer import score_snapshot
from recruit_crawler.schemas import JDSnapshot
from recruit_crawler.user_context import UserContextImportError, parse_context_document

CONFIG = ROOT / "config" / "sample_config.json"


class SampleUserContextConfigTests(unittest.TestCase):

    def test_sample_configs_do_not_embed_private_profile_data(self) -> None:
        for path in [ROOT / "config" / "sample_config.json", ROOT / "config" / "live_sources.sample.json"]:
            payload = path.read_text(encoding="utf-8")
            self.assertNotIn("PRIVATE_PROFILE_CANARY", payload)
            config = load_config(path, allow_real_sources=path.name == "live_sources.sample.json")
            self.assertEqual(config.profile.private_canaries, [])

class UserContextSchemaTests(unittest.TestCase):
    def test_config_profile_creates_user_context_contract(self) -> None:
        config = load_config(CONFIG)

        self.assertEqual(config.user_context.skills, config.profile.skills)
        self.assertEqual(config.user_context.explicit_deal_breakers, config.profile.exclusions)
        self.assertEqual(config.user_context.provenance["skills"], "config.profile.skills")

    def test_explicit_deal_breaker_excludes_without_raw_leakage(self) -> None:
        config = load_config(CONFIG)
        snapshot = JDSnapshot(
            source_id="fixture",
            source_url="https://jobs.example.test/unpaid",
            source_posting_id="unpaid",
            title="Unpaid Internship ML Engineer",
            company="Example",
            location="Seoul",
            deadline_raw=None,
            deadline=None,
            deadline_uncertain=False,
            required_qualifications=["Python", "machine learning"],
            preferred_qualifications=["LLM"],
            responsibilities=["Build models"],
            company_info=["AI team"],
        )

        assessment = score_snapshot(snapshot, config)

        self.assertEqual(assessment.verdict, "exclude")
        self.assertEqual(assessment.recommendation, "low_priority")
        self.assertEqual(assessment.deal_breaker_hits, ["unpaid internship"])
        self.assertNotIn("PRIVATE_PROFILE_CANARY", " ".join(assessment.risks))

class DocumentContextParserTests(unittest.TestCase):
    def test_plaintext_context_imports_user_context(self) -> None:
        context = parse_context_document(ROOT / "fixtures" / "user_context" / "context.md")

        self.assertIn("Python", context.skills)
        self.assertIn("Seoul", context.preferred_locations)
        self.assertEqual(context.max_experience_years, 2)

    def test_private_canary_document_fails_closed(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "private.md"
            path.write_text("Skills: Python\nPRIVATE_PROFILE_CANARY", encoding="utf-8")

            with self.assertRaises(UserContextImportError):
                parse_context_document(path)

    def test_docx_context_imports_user_context_when_dependency_available(self) -> None:
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "context.docx"
            doc = Document()
            doc.add_paragraph("Roles: ML Engineer")
            doc.add_paragraph("Skills: Python, SQL")
            doc.add_paragraph("Locations: Remote")
            doc.add_paragraph("Experience: 2 years")
            doc.save(str(path))

            context = parse_context_document(path)

        self.assertIn("SQL", context.skills)
        self.assertIn("Remote", context.preferred_locations)

    def test_pdf_context_imports_user_context_when_dependency_available(self) -> None:
        try:
            from pypdf import PdfWriter
        except ImportError:
            self.skipTest("pypdf is not installed")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "empty.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=72, height=72)
            with path.open("wb") as fh:
                writer.write(fh)

            with self.assertRaises(UserContextImportError):
                parse_context_document(path)
